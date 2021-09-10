#coding=utf-8
import sys
reload(sys)
sys.setdefaultencoding('utf-8')
from docx import Document
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #chinese
import os
from win32com import client as wc #导入doc转docx

#找到文件夹下的所有doxc文件并获得文件名
def file_name(file_dir):
    for root,dirs,files in os.walk(file_dir):
        return files

#遍历刚才找到的所有文件夹然后替换关键字
def change_text(old_text, new_text,document):
    all_paragraphs = document.paragraphs
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

    all_tables = document.tables
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

#把doc文件转为docx文件
def doc_to_docx(file1,file2):
    word = wc.Dispatch("Word.Application") # 打开word应用程序
    doc = word.Documents.Open(file1)       #打开word文件
    doc.SaveAs("{}x".format(file2), 12)    #另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close()                            #关闭原来word文件
    word.Quit()
    return "{}x".format(file)              #返回的是一个file对象或者文件名称

#把docx文件转为doc文件
def docx_to_doc(file1,file2):
    word = wc.Dispatch("Word.Application") # 打开word应用程序
    doc = word.Documents.Open(file1)       #打开word文件
    doc.SaveAs("{}".format(file2[:-1]), 0) #另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close()                            #关闭原来word文件
    word.Quit()
    return "{}x".format(file)  #返回的是一个file对象或者文件名称

#进行批量替换并生成新文件
#第一个参数传入模板文件夹路径，第二个参数导出结果到某个文件夹路径（请写绝对路径）
def deal_task(importPath,exportPath):
    words = file_name(importPath)   #获取所有文件的文件名
    for words_name in words:
        print words_name            #打印获取的文件名查看是否有错误
        #只有后缀为docx的文件可以继续执行
        if words_name.find('.docx') != -1:
            #PackageNotFoundError,捕获一下该异常
            try:
                document = Document(importPath+'/'+words_name)        #读取当前遍历的文件
                document.styles['Normal'].font.name = u"仿宋_GB2312"
                document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u"仿宋_GB2312")
                document.styles['Normal'].font.size = Pt(12)
                #进行替换操作
                #old_text和new_text都是动态传入
                #遍历entry_list和entry_str来执行动态操作，偶是old奇数是new
                for num in range(0,len(entry_list)-1,2):
                    #输入框任一一栏没有输入内容不会执行替换操作
                    if entry_list[num] != '' and entry_list[num+1] != '':
                        change_text(entry_list[num].strip(),entry_list[num+1].strip(),document)
                        print 'old_text: ' + entry_list[num]
                        print 'new_next: ' + entry_list[num+1]
                document.save(exportPath+"/"+"auto"+words_name)   #默认存储为auto+原文件名的方式
            except Exception as e:
                print '打开docx文件失败'
                print e
    print "完成！"

#进行批量转换并生成转换后的文件到文件夹内
#第一个参数导入的路径,第二个是导出的路径
def docTodocx(importPath,exportPath):
    fileNames = file_name(importPath)  # 获得所有文件名
    # filePath = importPath
    for fileName in fileNames:         # 遍历文件夹下的所有文件
        # 先判断是否有doc然后再判断docx，只要不是doc直接跳过
        # fileName = fileName.decode('gb2312').encode('utf-8')
        if fileName.find('doc') != -1:
            print '进入doc处理循环'
            file1 = importPath + "/" + fileName
            print 'file1:' + file1
            file1 = file1.decode('utf-8').encode('gb2312')
            # print 'gb2312编码的file1: ' + file1
            file2 = exportPath + "/" + fileName
            print 'file2:' + file2
            file2 = file2.decode('utf-8').encode('gb2312')
            doc_to_docx(file1, file2)

#设定自己要替换的数组第一个是old_text,第二个是new_text依次次类推
entry_list = ["X1","2021","X2","2022"]

if __name__ == "__main__":
    #如果模板全是docx只需要调用deal_task()函数就行了，不是的话先调用docTodocx转为docx

    #第一步把doc转为docx
    docTodocx("I:/pythonProject1/docx_operation/importPath","I:/pythonProject1/docx_operation/temporary")
    deal_task("I:/pythonProject1/docx_operation/temporary","I:/pythonProject1/docx_operation/exportPath")

